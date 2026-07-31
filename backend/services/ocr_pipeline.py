"""
ContractIQ — OCR Pipeline

Extracts text from PDF, DOCX, and TXT files.
Uses PyMuPDF for PDFs with EasyOCR fallback for scanned documents.
"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using PyMuPDF.
    Falls back to EasyOCR if no text is found (scanned PDF).
    """
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    text_parts = []

    for page_num, page in enumerate(doc):
        page_text = page.get_text("text")
        if page_text.strip():
            text_parts.append(page_text)

    doc.close()

    # If PyMuPDF found text, return it
    full_text = "\n\n".join(text_parts).strip()
    if full_text and len(full_text) > 50:
        logger.info(f"PyMuPDF extracted {len(full_text)} chars from {file_path}")
        return full_text

    # Fallback to EasyOCR for scanned PDFs
    logger.info(f"No text found via PyMuPDF, falling back to EasyOCR for {file_path}")
    return _ocr_fallback_pdf(file_path)


def _ocr_fallback_pdf(file_path: str) -> str:
    """Use EasyOCR as fallback for scanned PDFs."""
    try:
        import fitz
        import easyocr

        reader = easyocr.Reader(["en"], gpu=False)
        doc = fitz.open(file_path)
        text_parts = []

        for page_num, page in enumerate(doc):
            # Render page to image
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")

            # OCR the image
            import io
            from PIL import Image
            import numpy as np

            image = Image.open(io.BytesIO(img_bytes))
            img_array = np.array(image)

            results = reader.readtext(img_array, detail=0, paragraph=True)
            if results:
                text_parts.append("\n".join(results))

        doc.close()

        full_text = "\n\n".join(text_parts).strip()
        logger.info(f"EasyOCR extracted {len(full_text)} chars from {file_path}")
        return full_text

    except ImportError:
        logger.warning("EasyOCR not available, returning empty text")
        return ""
    except Exception as e:
        logger.error(f"EasyOCR failed: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document

    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    full_text = "\n\n".join(text_parts).strip()
    logger.info(f"Extracted {len(full_text)} chars from DOCX {file_path}")
    return full_text


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    encodings = ["utf-8", "latin-1", "cp1252"]

    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read().strip()
            logger.info(f"Extracted {len(text)} chars from TXT {file_path} ({encoding})")
            return text
        except UnicodeDecodeError:
            continue

    logger.error(f"Could not decode TXT file {file_path}")
    return ""


def extract_text(file_path: str, file_type: str) -> str:
    """
    Main entry point: extract text from any supported file type.

    Args:
        file_path: Path to the file
        file_type: File extension without dot (pdf, docx, txt)

    Returns:
        Extracted text content
    """
    file_type = file_type.lower().lstrip(".")

    extractors = {
        "pdf": extract_text_from_pdf,
        "docx": extract_text_from_docx,
        "txt": extract_text_from_txt,
    }

    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")

    text = extractor(file_path)

    if not text:
        logger.warning(f"No text extracted from {file_path}")

    return text
